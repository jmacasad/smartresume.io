# Flask backend for resume analysis and tailoring
from flask import Flask, request, jsonify
from flask_cors import CORS
import requests
import base64
import tempfile
import uuid
from bs4 import BeautifulSoup
from docx import Document
from nltk.corpus import wordnet
from fuzzywuzzy import fuzz
import re
import hashlib
import os
import json
from utils.parser import decode_resume_base64, extract_experience_block
import docx2txt
from docx import Document
from rapidfuzz import fuzz
from resume_builder import generate_hybrid_resume
from utils.prompt_builder import generate_hybrid_resume_prompt
from utils.section_confidence import flag_low_confidence_sections
from utils.diff_check import compare_resumes

import nltk
from nltk.stem import PorterStemmer

def decode_resume_base64(resume_text):
    """Helper to decode a base64-encoded .docx file or plain text."""
    if resume_text.startswith('UEsDB'):
        try:
            docx_bytes = base64.b64decode(resume_text)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                f.write(docx_bytes)
                path = f.name
            doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            raise ValueError(f"Resume decode failed: {e}")
    else:
        return resume_text

import pdfkit

PDF_CONFIG = pdfkit.configuration(wkhtmltopdf='/usr/local/bin/wkhtmltopdf')

from hybrid_template import generate_hybrid_resume


def extract_job_title_from_text(jd_text):
    lines = jd_text.strip().split('\n')[:15]

    # 🧠 Try Regex Match First
    for line in lines:
        match = re.search(
            r'seeking an experienced (.+?)( to|\.)', line.lower())
        if match:
            title = match.group(1).strip().title()
            return title

    # 🛡️ Fallback – Extract line before dash (like "X – Location")
    for line in lines:
        clean = line.strip()
        if "–" in clean:
            title_candidate = clean.split("–")[0].strip()
            if 3 <= len(title_candidate.split()) <= 8:
                return title_candidate

    # 🪂 Final fallback – any uppercase line that doesn't sound generic
    for line in lines:
        clean = line.strip()
        if clean and clean[0].isupper() and not clean.lower().startswith((
            "about", "join", "responsibilities", "we’re", "we are", "qualifications"
        )):
            return clean

    return "Unknown Role"

stemmer = PorterStemmer()

def normalize_skills(text: str) -> list[str]:
    parts = re.split(r'[\n,;•]+', text)
    cleaned = set()
    for p in parts:
        p = re.sub(r'^\d+\.?\s*', '', p)
        p = re.sub(r'(key (hard|soft) skills:|additional qualifications:|eg|required skills:)', '', p, flags=re.I)
        p = re.sub(r'[^a-zA-Z0-9\s\-]', '', p).strip().lower()

        if p:
            # Handle hyphenated words (treat them as single units)
            if '-' in p:
                cleaned.add(p)
            else:
                # Split into individual words and stem them
                words = p.split()
                stemmed_words = [stemmer.stem(word) for word in words]
                cleaned.add(" ".join(stemmed_words))

    return sorted(cleaned)

def extract_normalized_resume_skills(resume_text):
    """
    Extract and normalize skills from resume using the same logic as job description.
    """
    import re
    from nltk.stem import PorterStemmer
    stemmer = PorterStemmer()
    parts = re.split(r'[\n,;•]+', resume_text)
    cleaned = set()
    for p in parts:
        p = re.sub(r'^\d+\.?\s*', '', p)
        p = re.sub(r'(key (hard|soft) skills:|additional qualifications:|eg|required skills:)', '', p, flags=re.I)
        p = re.sub(r'[^a-zA-Z0-9\s\-]', '', p).strip().lower()
        if p:
            if '-' in p:
                cleaned.add(p)
            else:
                words = p.split()
                stemmed_words = [stemmer.stem(word) for word in words]
                cleaned.add(" ".join(stemmed_words))
    return sorted(cleaned)

def get_cache_key(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def load_cached_skills(job_text):
    os.makedirs("cache", exist_ok=True)
    cache_key = get_cache_key(job_text)
    cache_file = f"cache/{cache_key}.json"
    if os.path.exists(cache_file):
        with open(cache_file, "r") as f:
            return json.load(f)
    return 


def save_cached_skills(job_text, data):
    cache_key = get_cache_key(job_text)
    cache_file = f"cache/{cache_key}.json"
    with open(cache_file, "w") as f:
        json.dump(data, f, indent=2)
        
def generate_session_id():
    return str(uuid.uuid4())

def save_to_session(session_id, filename, content):
    folder = os.path.join("sessions", session_id)
    os.makedirs(folder, exist_ok=True)
    with open(os.path.join(folder, filename), "w", encoding="utf-8") as f:
        f.write(content)
    return os.path.join(folder, filename)

def extract_experience_block(resume_text):
    """Extracts the experience blocks from the resume text."""

    lines = resume_text.splitlines()
    start_index = -1
    for i, line in enumerate(lines):
        if "PROFESSIONAL EXPERIENCE" in line.upper():
            start_index = i
            break

    if start_index == -1:
        return []

    experience_text_lines = lines[start_index + 1:]  # Lines after the heading
    blocks = []
    current_block_lines = []
    in_experience = False

    for line in experience_text_lines:
        line = line.strip()
        if not line:
            continue  # Skip empty lines

        # Check for a line that looks like a title (contains "|" and is not "Referee details")
        if "|" in line and "Referee details" not in line:
            if current_block_lines:
                blocks.append("\n".join(current_block_lines))
            current_block_lines = [line]
            in_experience = True
        elif in_experience:
            current_block_lines.append(line)

    if current_block_lines:
        blocks.append("\n".join(current_block_lines))

    # Further refine blocks to ensure they have bullet points
    refined_blocks = []
    for block in blocks:
        if any(line.startswith("•") or line.startswith("-") for line in block.splitlines()):
            refined_blocks.append(block)

    return refined_blocks

def get_synonyms(word):
    return [word] # Only return the original word

def match_skills(skills_required, resume_text, scorer=fuzz.token_set_ratio, threshold=85, structured_experience=None):
    matched = []
    missing = []
    scores = {}
    raw_scores = {}

    for skill in skills_required:
        score = scorer(skill.lower(), resume_text.lower())
        scores[skill] = score
        raw_scores[skill] = score
        if score >= threshold:
            matched.append(skill)
        else:
            missing.append(skill)

    return matched, missing, scores, raw_scores

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

def extract_text_from_docx(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

OPENROUTER_API_KEY = os.getenv ("OPENROUTER_API_KEY") or "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"

def call_your_model_api(prompt):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }

    # 🔍 Step 1: Detect if prompt is a preformatted list of messages (e.g., from hybrid_prompt)
    if isinstance(prompt, list) and all(isinstance(msg, dict) and "role" in msg and "content" in msg for msg in prompt):
        messages = prompt
    else:
        messages = [
            {"role": "user", "content": prompt}
        ]

    # 🔧 Step 2: Build payload dynamically
    payload = {
        "model": "openai/gpt-4-turbo",
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 3000
    }

    # 🧠 Optional: Print preview of final messages going to LLM
    print("🧠 Final LLM prompt payload:", json.dumps(messages, indent=2))

    # 🚀 Step 3: Send to OpenRouter
    response = requests.post("https://openrouter.ai/api/v1/chat/completions", json=payload, headers=headers)
    response.raise_for_status()

    content = response.json()["choices"][0]["message"]["content"]
    return { "content": content }



jd_skill_cache = {}


@app.route('/analyze-jd-vs-resume', methods=['POST'])
def analyze_jd_vs_resume():
    global re  # Add this line at the beginning of the function
    data = request.json
    job_text = data.get("jobText", "").strip()
    resume_raw = data.get("resumeText", "").strip()

    # 🧠 NEW: Auto-detect job title and prepend it to the JD
    job_title = extract_job_title_from_text(job_text)

    print("🧾 Raw Job Description Received:", job_text[:500])
    print("🎯 Extracted Job Title:", job_title)

    job_text = f"Job Title: {job_title}\n\n{job_text}"

    if not job_text or not resume_raw:
        return jsonify({'error': 'Job description and resume text are required'}), 400

    # 🪓 Limit the size of job_description and resume to avoid token overflow
    job_text = job_text[:2500]
    resume_input = resume_raw[:2500]

    match_score = 0
    matched = []
    missing = []
    scores = {}       # Initialize scores, raw_scores, matched, missing here
    raw_scores = {}

    # 🧠 Step 1: Decode resume
    try:
        resume_input = decode_resume_base64(resume_raw)
    except Exception as e:
        return jsonify({"error": f"Resume parsing failed: {e}"}), 500

    session_id = generate_session_id()
    save_to_session(session_id, "jd.txt", job_text)
    save_to_session(session_id, "resume_raw.txt", resume_raw)
    save_to_session(session_id, "resume_decoded.txt", resume_input)

    # Extract experience blocks
    experience_blocks = extract_experience_block(resume_input)
    save_to_session(session_id, "resume_experience_blocks.txt", "\n\n".join(experience_blocks))

    # Structure the experience data
    structured_experience = []
    for block in experience_blocks:
        # Refined bullet point splitting
        lines = block.strip().split('\n')
        title = lines[0] if lines else "Unknown Title"
        descriptions = []
        for line in lines[1:]:
            line = line.strip()
            if line and (line.startswith('•') or line.startswith('-')):
                descriptions.append(line[1:].strip())  # Remove the bullet point character
            elif line:
                # Handle cases where bullet points might be missing or inconsistent
                if descriptions:
                    descriptions[-1] += " " + line  # Append to the previous description
                else:
                    descriptions.append(line)
        structured_experience.append({"title": title, "descriptions": descriptions})

    save_to_session(session_id, "resume_experience_structured.json", json.dumps(structured_experience, indent=2))  # Save structured data

    # 🔍 Extract skills prompt must be BEFORE try block!
    extract_skills_prompt = [
        {"role": "system", "content": "You are an expert at analyzing job descriptions to extract key skills. Identify both hard and soft skills explicitly mentioned or strongly implied by the responsibilities and qualifications. If possible, categorize them as 'essential' or 'desirable'."},
        {"role": "user", "content": f"Analyze the following job description and list the key skills, noting if they appear to be essential or desirable:\n{job_text}"}
    ]

    try:
        cache_dir = "cache"
        os.makedirs(cache_dir, exist_ok=True)

        # 🧠 Hash job_text to create a unique cache key
        cache_key = hashlib.md5(job_text.encode("utf-8")).hexdigest()
        cache_path = os.path.join(cache_dir, f"skills_{cache_key}.json")

        if os.path.exists(cache_path):
            print(f"📂 Using cached skills: {cache_path}")
            with open(cache_path, "r") as f:
                skill_data = json.load(f)
        else:
            print(f"🧠 Fetching skills from LLM...")
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            }
            skill_payload = {
                "model": "openai/gpt-4-turbo",
                "temperature": 0,
                "max_tokens": 500,
                "messages": extract_skills_prompt
            }
            skill_response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions", headers=headers, json=skill_payload)
            skill_data = skill_response.json()

            with open(cache_path, "w") as f:
                json.dump(skill_data, f, indent=2)

        print("🔍 OpenRouter skill response:", skill_data)

        if "choices" not in skill_data:
            return jsonify({'error': 'Skill extraction failed: no choices returned from LLM'}), 500

        skill_list_raw = skill_data['choices'][0]['message']['content']

        import re

        # 🧠 Split on comma OR newline — handles both lists and numbered formats
        raw_skills = re.split(r'[\n,]+', skill_list_raw)

        # 🧼 Normalize and clean each skill
        extracted_skills = []
        for skill in raw_skills:
            clean = skill.lower()
            clean = re.sub(r'^\d+\.?\s*', '', clean)
            clean = re.sub(
                r'(key (hard|soft) skills:|additional qualifications:|qualifications:|preferred qualifications:|required skills:)', '', clean)
            clean = re.sub(r'[^a-zA-Z0-9\s\-]', '', clean).strip()
            if clean:
                extracted_skills.append(clean)

        # 🧠 Deduplicate
        extracted_skills = list(set(extracted_skills))

        # 🧠 Normalize resume using JD-style logic
        normalized_resume_skills = extract_normalized_resume_skills(resume_input)
        resume_input_clean = " ".join(normalized_resume_skills)

        print("🧾 Cleaned Resume Text Preview (first 1000 chars):")
        print(resume_input_clean[:1000])
        print("🧠 Extracted Skills to Match:")
        print(extracted_skills)

        matched, missing, scores, raw_scores = match_skills(
            extracted_skills,
            resume_input_clean,
            scorer=fuzz.partial_ratio,
            threshold=60,
            structured_experience=structured_experience # Pass structured data
        )

        print("🔍 Skill Match Scores:", scores)
        print("🔍 Raw Skill Match Scores:", raw_scores)
        print("🔍 Normalized Resume Skills:", normalized_resume_skills)
        print("🔍 Resume Clean Tokens (for match):")

        # (Optional) debug print all scores:
        for skill, sc in scores.items():
            print(f"🔍 Skill Match Check -> {skill} :: Score: {sc}")

        match_score = round((len(matched) / len(extracted_skills)) * 100, 2) if extracted_skills else 0


        print("📊 Skill Stats")
        print("• Extracted from JD:", len(extracted_skills))
        print("• Matched in Resume:", len(matched))
        print("• Missing from Resume:", len(missing))
        print("🧠 Skills Extracted:", extracted_skills)
        print("✅ Match Score:", match_score)

    except Exception as e:
        print("⚠️ Skill extraction failed:", str(e))
        match_score = 0
        matched = []
        missing = []

    # Step: Generate alignment summary and follow-up questions
    job_text_trunc = job_text[:4000]
    resume_input_trunc = resume_input[:4000]

    # Hard truncate all big chunks to stay under 16K tokens
    job_text_trunc = job_text[:3000]
    resume_input_trunc = resume_input[:3000]

    print("🧾 Final job_text sent to GPT:\n", job_text[:1000])
    print("📄 Resume snippet sent to GPT:\n", resume_input[:1000])

    prompt = [
        {"role": "user", "content": f'''
You are helping a job seeker improve their resume to match a specific job description.

Job Title: {job_title}

Job Description:
{job_text_trunc}

Candidate Resume:
{resume_input_trunc}

1.  Write a clear and concise summary (5–8 full sentences, paragraph format) explaining how well the candidate’s resume aligns with the job. Mention strengths and areas of alignment, and clearly state any weaknesses or gaps.

2.  Then, write at least 5 clarification questions for the **job seeker** — not the recruiter. These should help improve their resume by addressing missing skills, unclear responsibilities, or experience gaps based on the job description.

❌ Do NOT ask questions that can be answered by reading the job description itself.
❌ Do NOT ask general HR questions (e.g. “Does the JD mention...?”)
❌ Only ask questions that help improve the resume for a better ATS score.

Return your answer in this exact format:

[Paragraph summary goes here]

Clarification Questions:
1.  ...
2.  ...
3.  ...
'''}
    ]

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "openai/gpt-4-turbo",
        "messages": prompt
    }

    try:
        ai_response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
        if ai_response.status_code != 200:
            return jsonify({'error': f'OpenRouter error: {ai_response.text}'}), ai_response.status_code

        ai_data = ai_response.json()

        # 👁️ Debug print
        print("🧠 Claude-style AI response (summary+questions):", ai_data)

        if "choices" not in ai_data:
            return jsonify({'error': 'Failed to generate summary & questions: missing choices from model'}), 500

        analysis_output = ai_data['choices'][0]['message']['content']

        # Split output into summary and follow-up questions
        lines = analysis_output.strip().split("\n")
        summary_lines = []
        questions = []
        in_questions = False

        for line in lines:
            if "Clarification Questions" in line:
                in_questions = True
                continue
            if in_questions:
                if line.strip().startswith("*") or line.strip().startswith("-") or line.strip()[0:1].isdigit():
                    question = line.strip().lstrip("*-1234567890. ").strip()
                    if question:
                        questions.append(question)
            else:
                summary_lines.append(line.strip())

        summary = "\n".join(summary_lines).strip()

        print("Questions parsed:", questions)

        print("📤 Returning ATS Analytics:", match_score, matched, missing)

        # ✅ Send back structured summary + questions
        return jsonify({
            "summary": summary,
            "questions": questions,
            "analytics": {
                "match_score": match_score,
                "matched_skills": matched,
                "missing_skills": missing,
                "skill_scores": scores,
                "raw_skills_scores": raw_scores,
                "structured_experience": structured_experience  # Send structured data
            },
            "session_id": session_id
        })

    except Exception as e:
        return jsonify({'error': f'Failed to analyze: {str(e)}'}), 500
    
# 🧠 Helper: Construct hybrid prompt with JD, Resume, and extra Q&A
def generate_hybrid_resume(context):
    job_text = context.get("job_text", "").strip()
    resume_raw = context.get("resume_raw", "").strip()
    extra_answers = context.get("extra_answers", [])

    # ✅ Detect job title
    job_title = extract_job_title_from_text(job_text)
    job_text = f"Job Title: {job_title}\n\n{job_text}"

    # ✅ Compile Q&A into readable format
    qna_context = ""
    if extra_answers:
        qna_context += "Additional context provided by the candidate:\n"
        for i, answer in enumerate(extra_answers, 1):
            qna_context += f"{i}. {answer.strip()}\n"

    # 🧠 Final Prompt Construction
    prompt = f"""
You are a resume optimization assistant. Based on the job description below, optimize the resume to:
- Match the key hard and soft skills
- Use ATS-friendly formatting and keywords
- Expand each experience to include 6+ detailed bullet points
- Add a section for 'Project Key Achievements' relevant to the JD

### Job Description:
{job_text}

### Candidate Resume:
{resume_raw}

### Clarifying Context:
{qna_context}

Return only the optimized resume text.
"""
    return prompt.strip()



@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    data = request.json
    resume_text = data.get("resumeText", "")
    job_text = data.get("jobText", "")
    extra_answers = data.get("extraAnswers", [])
    session_id = data.get("sessionId") or generate_session_id()
    extracted_skills = data.get("skills", [])
    
    def clean_skill(skill):
        return re.split(r'\b(such as|including|experience|familiarity|ability to|proficiency|knowledge|expertise)\b', skill, flags=re.I)[0].strip()
    extracted_skills = [clean_skill(s) for s in extracted_skills if isinstance(s, str) and s.strip()]
    skills_csv = ", ".join(extracted_skills)
    skills_csv = ", ".join(extracted_skills)
    
    resume_text = decode_resume_base64(resume_text)

    parsed = {
        "resume_text": resume_text,
        "job_text": job_text,
        "extra_answers": extra_answers,
        "experience": extract_experience_block(resume_text)
    }

    # 🚨 Confidence check
    confidence_flags = flag_low_confidence_sections(parsed)

    # ✅ Add extracted skills into parsed resume dictionary
    parsed["skills"] = extracted_skills

    # 🧠 Build hybrid prompt
    hybrid_prompt = generate_hybrid_resume(parsed)
    print("🧠 Parsed Resume Structure:", parsed)
    print("🧠 Injected Skills CSV:", ", ".join(extracted_skills))
    
    # 🧠 Final Prompt Preview Debug
    print("🧠 Final Prompt Preview:\n")

    if isinstance(hybrid_prompt, list) and all(isinstance(msg, dict) for msg in hybrid_prompt):
        for msg in hybrid_prompt:
            print(f"{msg['role'].upper()}:\n{msg['content']}\n")
    else:
        print("⚠️ hybrid_prompt is not in expected format. Here's what we received:")
        print(hybrid_prompt)


    # 🧠 Call the model
    llm_response = call_your_model_api(hybrid_prompt)
    resume_output = llm_response['content']
    
    # ⚠️ Strip code block formatting from GPT output
    import re
    resume_output = re.sub(r"^```(?:\w+)?\n?", "", resume_output.strip())
    resume_output = re.sub(r"\n?```$", "", resume_output.strip())

    
    # 🧠 Clean the resume output
    clean_text = BeautifulSoup(resume_output, "html.parser").get_text().lower()
    clean_text = re.sub(r'[^a-z0-9\s\-]', '', clean_text)
    clean_text = re.sub(r'\s+', ' ', clean_text).strip()

    # ---- 3) Load or fetch & normalize skills ----
    skill_data = load_cached_skills(job_text)
    if not skill_data:
        skill_payload = {
            "model": "openai/gpt-4-turbo",
            "temperature": 0,
            "max_tokens": 300,
            "messages": [
                {"role": "system", "content": "Extract key hard and soft skills from job description."},
                {"role": "user",   "content": job_text}
            ]
        }
        skill_headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }
        skill_res = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=skill_headers,
            json=skill_payload
        )
        skill_data = skill_res.json()
        save_cached_skills(job_text, skill_data)

    if "choices" not in skill_data:
        return jsonify({'error': 'Skill extraction failed'}), 500

    skill_text = skill_data['choices'][0]['message']['content']
    extracted_skills = normalize_skills(skill_text)
    if len(extracted_skills) > 25:
        extracted_skills = extracted_skills[:25]
        
    # ✅ Add extracted skills into parsed resume dictionary
    parsed["skills"] = extracted_skills

    # 🧠 Optional but recommended: Debug log
    print("🧠 Injected Skills CSV:", ", ".join(extracted_skills))

    # Match extracted JD skills
    matched, missing, scores = match_skills(extracted_skills, clean_text, scorer=fuzz.partial_ratio, threshold=75)
    match_score = round((len(matched) / len(extracted_skills)) * 100, 2) if extracted_skills else 0

    # ---- Preprocess resume text ----
    if not resume_text:
        return jsonify({'error': 'No resume text provided'}), 400

    # Treat input as plain text (React sends text not docx)
    resume_input = resume_text.strip()
    resume_input_trunc = resume_input[:3000]

    # ---- 5) Generate ATS-optimized resume, score & summary ----
    try:
        parsed_resume_data = {
            "resume_text": resume_input_trunc,
            "job_text": job_text,
            "extra_answers": extra_answers,
            "experience": extract_experience_block(resume_input_trunc)
        }

        # Create prompt
        prompt = f"""
            You are an expert resume optimizer.

            Instructions:
            - Do NOT create fake jobs, companies, dates, or degrees.
            - ONLY use real experience from the candidate's original resume.
            - Emphasize any overlapping experience with the job description.
            - If required data is missing, leave it out — do not guess.
            - Use the skill list exactly as written.

            Job Description:
            {job_text.strip()}

            Resume Provided (truncated):
            {resume_input_trunc.strip()}

            Required Skills:
            {skills_csv}

            Additional Info:
            {chr(10).join(extra_answers)}

            Your task:
            - Reformat the resume to match the JD.
            - Highlight relevant achievements.
            - Make it ATS-friendly and markdown-structured.

            Return ONLY the optimized resume.
            """


        # Make API call
        llm_response = call_your_model_api(prompt=prompt)
        resume_output = llm_response['content']
        
        # ⚠️ Strip code block formatting from GPT output
        resume_output = re.sub(r"^```(?:\w+)?\n?", "", resume_output.strip())
        resume_output = re.sub(r"\n?```$", "", resume_output.strip())

        
        # 🔍 Show original experience and updated resume output side-by-side
        print("\n🆚 BEFORE: Extracted Experience\n")
        print(parsed_resume_data.get("experience", "")[:1500])
        
        # 🧠 NEW: Run diff analysis between original and rewritten resume
        diff_report = compare_resumes(
            original=parsed.get("resume_text", ""),
            rewritten=resume_output,
            min_bullets_per_job=6
        )

        # 🧠 Debug print to console
        print("\n🧠 Resume Quality Check:")
        print("🟡 Skills Dropped:", diff_report.get("missing_skills"))
        print("🔴 Bullet Gaps:", diff_report.get("bullet_issues"))
        print("📄 Section Diffs:\n", diff_report.get("text_diffs"))


        print("\n✅ AFTER: Optimized Resume Output (First 1500 characters)\n")
        print(resume_output[:1500])
        
        session_id = request.json.get("sessionId") or generate_session_id()
        
        save_to_session(session_id, "resume_optimized.txt", resume_output)
                
        print("\n🔬 Resume Output Preview (first 1000 chars):\n", resume_output[:1000])
        print("\n🧠 Skills CSV Sent to AI:\n", skills_csv)

        # 5b) Clean & fuzzy‑match for analytics
        soup = BeautifulSoup(resume_output, "html.parser")
        clean_text = soup.get_text().lower()
        clean_text = re.sub(r'[^a-z0-9\s\-]', '', clean_text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()

        # ✅ Store and return results
        save_to_session(session_id, "ats_final.json", json.dumps({
            "match_score": match_score,
            "matched_skills": matched,
            "missing_skills": missing
        }, indent=2))
        
        matched, missing, scores = match_skills(
            extracted_skills,
            clean_text,
            scorer=fuzz.partial_ratio,
            threshold=75
        )
        match_score = round((len(matched) / len(extracted_skills)) * 100, 2) if extracted_skills else 0
        
        print(f"🧪 match_score: {match_score}")
        print(f"🧪 matched_skills: {matched}")
        print(f"🧪 missing_skills: {missing}")

        return jsonify({
            "resume": resume_output,
            "analytics": {
                "match_score": match_score,
                "matched_skills": matched,
                "missing_skills": missing
            },
            "session_id": session_id,
            "confidence_flags": confidence_flags
        })
        

    except Exception as e:
        return jsonify({'error': f'Failed to generate resume: {e}'}), 500

@app.route('/download-resume/<session_id>', methods=['GET'])
def download_resume(session_id):
    folder = os.path.join("sessions", session_id)
    resume_path = os.path.join(folder, "resume_optimized.txt")
    pdf_path = os.path.join(folder, "resume_final.pdf")

    if not os.path.exists(resume_path):
        return jsonify({"error": "Resume not found"}), 404

    with open(resume_path, "r", encoding="utf-8") as f:
        raw_resume = f.read()

    html = f"""
    <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; padding: 30px; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; }}
            </style>
        </head>
        <body>
            <pre>{raw_resume}</pre>
        </body>
    </html>
    """

    pdfkit.from_string(html, pdf_path, configuration=PDF_CONFIG)
    return send_file(pdf_path, as_attachment=True, download_name="optimized_resume.pdf")


if __name__ == '__main__':
    job_description_skills = """
    - Project Management Skills
    - Managing Projects
    - Agile Methodologies
    - Scrum
    - Communication skills
    - Excellent Communication
    - Problem-solving abilities
    - Problem Solving
    """
    normalized_jd_skills = normalize_skills(job_description_skills)
    print("Normalized JD Skills:", normalized_jd_skills)

    resume_skills_text = """
    Experience in managing project timelines
    Proficient in Agile and Scrum frameworks
    Strong communication ability
    Good at problem-solving
    """
    normalized_resume_skills = extract_normalized_resume_skills(resume_skills_text)
    print("Normalized Resume Skills:", normalized_resume_skills)

    app.run(debug=True)     
