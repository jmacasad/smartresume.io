# Flask backend for resume analysis and tailoring
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
import base64
import tempfile
import uuid
from bs4 import BeautifulSoup
from docx import Document
from nltk.corpus import wordnet
from fuzzywuzzy import fuzz
from routes.analyze import analyze_bp
import re
import hashlib
import os
import traceback
import json
from utils.parser import decode_resume_base64, extract_experience_block
import docx2txt
from docx import Document
from rapidfuzz import fuzz
from resume_builder import generate_hybrid_resume
from utils.prompt_builder import generate_hybrid_resume_prompt
from utils.section_confidence import flag_low_confidence_sections
from utils.diff_check import compare_resumes
from utils.parser import decode_resume_base64, extract_experience_block
import skill_normalizer
import markdown2

import nltk
# List of NLTK resources needed
nltk_resources = [
    ('corpora/wordnet', 'wordnet'),
    ('tokenizers/punkt', 'punkt'),
    ('corpora/stopwords', 'stopwords')
]

print("--- Checking NLTK resources ---")
for resource_path, resource_name in nltk_resources:
    try:
        nltk.data.find(resource_path)
        print(f"Resource '{resource_name}' found.")
    except LookupError:
        print(f"Resource '{resource_name}' not found. Attempting download...")
        try:
            nltk.download(resource_name, quiet=True) # Use quiet=True for cleaner logs
            print(f"Resource '{resource_name}' downloaded successfully.")
            # Optional: You could try nltk.data.find again here to confirm
        except Exception as e:
            # Handle potential download errors (network issues, etc.)
            print(f"!!! ERROR: Failed to download NLTK resource '{resource_name}': {e} !!!")
            print("!!! Please check your internet connection or try downloading manually: python -m nltk.downloader <resource_name> !!!")
            # Decide if you want to exit or continue with potentially limited functionality
            # exit(1) # Uncomment to stop the app if download fails

print("--- NLTK resource check complete ---")

    
from nltk.stem import PorterStemmer
import pdfkit

#---CONFIGURATION---
OPENROUTER_API_KEY = os.getenv ("OPENROUTER_API_KEY") or "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"

#--Path configuration--
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SESSIONS_BASE_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "sessions")

print(f"--- INFO: Script directory: {SCRIPT_DIR} ---")
print(f"--- INFO: Sessions base directory set to: {SESSIONS_BASE_DIR} ---")
print(f"---      (Ensure this directory exists and the Flask process has write permissions) ---")

PDF_CONFIG_PATH = '/usr/bin/wkhtmltopdf'
if not os.path.exists(PDF_CONFIG_PATH):
     print(f"--- WARNING: wkhtmltopdf not found at {PDF_CONFIG_PATH}. PDF download may fail. ---")
     # Add paths for other OS if necessary, e.g., for Windows or macOS (brew install)
     # PDF_CONFIG_PATH_MAC = '/usr/local/bin/wkhtmltopdf'
     # if os.path.exists(PDF_CONFIG_PATH_MAC): PDF_CONFIG_PATH = PDF_CONFIG_PATH_MAC
     PDF_CONFIG = None # Or handle gracefully later
else:
    PDF_CONFIG = pdfkit.configuration(wkhtmltopdf=PDF_CONFIG_PATH)
    
    
#--- Helper functions ---

stemmer = PorterStemmer()

def decode_resume_base64(resume_text):
    """Helper to decode a base64-encoded .docx file or plain text."""
    if isinstance(resume_text, str) and resume_text.startswith(('UEsDB', 'PK')):
        try:
            # Ensure padding is correct for base64
            missing_padding = len(resume_text) % 4
            if missing_padding:
                resume_text += '=' * (4 - missing_padding)
            docx_bytes = base64.b64decode(resume_text, validate=True)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_f:
                temp_f.write(docx_bytes)
                temp_docx_path = temp_f.name
            # Use python-docx to extract text
            doc = Document(temp_docx_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            os.remove(temp_docx_path) # Clean up temp file
            if not full_text: # Fallback if docx reading fails somehow
                 print("--- WARNING: python-docx extracted empty text, trying docx2txt ---")
                 try:
                      full_text = docx2txt.process(temp_docx_path) # Requires docx2txt
                 except Exception as e2:
                      print(f"--- ERROR: docx2txt fallback failed: {e2} ---")
                      raise ValueError(f"Resume DOCX decode failed with both libraries.")
            return full_text
        except (base64.binascii.Error, FileNotFoundError, Exception) as e:
            print(f"--- ERROR: Resume decode failed: {e} ---")
            # Fallback: treat as plain text if decoding fails badly
            return resume_text
        finally:
            # Ensure cleanup even if errors occur during processing
            if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
                try:
                    os.remove(temp_docx_path)
                except OSError:
                    pass # Ignore cleanup error
    elif isinstance(resume_text, str):
         # Assume it's plain text if not starting with DOCX identifier
         return resume_text
    else:
         # Handle unexpected types
         print(f"--- WARNING: Unexpected type for resume_text: {type(resume_text)}. Returning empty string. ---")
         return ""


def extract_job_title_from_text(jd_text):
    lines = jd_text.strip().split('\n')[:15]

    # 🧠 Try Regex Match First
    for line in lines:
        match = re.search(
            r'seeking an experienced (.+?)( to|\.)', line.lower())
        if match:
            title = match.group(1).strip().title()
            return title

    # 🛡️ Fallback – Extract line before dash (like "X – Location")
    for line in lines:
        clean = line.strip()
        if "–" in clean:
            title_candidate = clean.split("–")[0].strip()
            if 3 <= len(title_candidate.split()) <= 8:
                return title_candidate

    # 🪂 Final fallback – any uppercase line that doesn't sound generic
    for line in lines:
        clean = line.strip()
        if clean and clean[0].isupper() and not clean.lower().startswith((
            "about", "join", "responsibilities", "we’re", "we are", "qualifications"
        )):
            return clean

    return "Unknown Role"

stemmer = PorterStemmer()
 
def normalize_skills(text: str) -> list[str]:
  """Extracts normalized skills from a resume using an AI agent (LLM)."""
  parts = re.split(r'[\n,;•]+', text)
  cleaned = set()
  for p in parts:
    p = re.sub(r'^\d+\.?\s*', '', p)
    p = re.sub(r'(?:key |hard |soft )?skills?:?', '', p, flags=re.I)
    p = re.sub(r'additional qualifications:?', '', p, flags=re.I)
    p = re.sub(r'e\.?g\.?:?', '', p, flags=re.I)
    p = re.sub(r'required skills:?', '', p, flags=re.I)
    p = re.sub(r'preferred qualifications:?', '', p, flags=re.I)
    p = re.sub(r'\s\(.*\)', '', p)
    p = re.sub(r'[^a-zA-Z0-9\s\-]', '', p).strip().lower()
 

  if p and len(p) > 1:
  # Call LLM to normalize the skill
    prompt = [
  {"role": "system", "content": "You are an expert in resume parsing and skill normalization."},
  {"role": "user", "content": f"""
  Extract key skills and qualifications from the resume text below. A skill can be a technical skill, a soft skill, a tool, a technology, or a qualification.  Return only the skill phrases, separated by commas.  If no skills are found, return an empty string.
  
  Resume Text:
  {text}
  
  Output the extracted skills as a comma-separated list.
  
  Examples:
  
  Input:  "Experience with Python and Java. Strong communication and problem-solving skills.  Certified Project Manager."
  Output: Python, Java, communication, problem solving, Certified Project Manager
  
  Input: "Familiar with  Microsoft 365, SharePoint, and  Azure AD.  Agile Methodologies,  Budget Management."
  Output: Microsoft 365, SharePoint, Azure AD, Agile Methodologies, Budget Management
  
  Input: "No skills mentioned."
  Output: 
  """}
  ]
    try:
        response_data = call_your_model_api(prompt)
        if response_data and response_data['content']:
            skills_text = response_data['content']
            normalized_skills = [skill.strip() for skill in skills_text.split(',')]
            cleaned.update(normalized_skills)
        else:
            print(f"--- WARNING: LLM returned empty response for skill normalization.  Proceeding with original text. ---")
            cleaned.add(p)
    except Exception as e:
        print(f"!!! ERROR: LLM call failed during skill normalization: {e}.  Using original skill. !!!")
        traceback.print_exc()
        cleaned.add(p)  # Keep the original skill if LLM call fails
 

  generic_terms = {"communication", "teamwork", "problem solving",
  "detail oriented", "etc", "and", "or", "including",
  "strong", "excellent", "good", "ability", "proficient"}
  cleaned = {s for s in cleaned if s not in generic_terms and not s.isdigit()}
  return sorted(list(cleaned))


def extract_normalized_resume_skills(resume_text):
    """Extracts normalized skills from a resume using an AI agent (LLM)."""
    prompt = [
        {"role": "system",
         "content": "You are an expert resume analyst. Your task is to identify and extract skills from the provided resume text."},
        {"role": "user", "content": f"""
        Identify and extract all skills from the resume text below.  A skill can be a technical skill, a soft skill, a tool, a technology, or a qualification.  Extract only the skill phrases.
        
        Resume Text:
        {resume_text}
        
        Output the extracted skills as a comma-separated list.
        
        Examples:
        
        Input:  "Experience with Python and Java. Strong communication and problem-solving skills.  Certified Project Manager."
        Output: Python, Java, communication, problem solving, Certified Project Manager
        
        Input: "Familiar with  Microsoft 365, SharePoint, and  Azure AD.  Agile Methodologies,  Budget Management."
        Output: Microsoft 365, SharePoint, Azure AD, Agile Methodologies, Budget Management
        """}
    ]
    try:
        response_data = call_your_model_api(prompt)
        skills_text = response_data['content']
        normalized_skills = [skill.strip() for skill in skills_text.split(',')]
        return normalized_skills
    except Exception as e:
        print(f"!!! ERROR: Skill extraction from resume failed: {e} !!!")
        traceback.print_exc()
        return []  # Return an empty list on error to avoid crashing

def get_cache_key(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def load_cached_skills(job_text):
    cache_dir = "cache"
    os.makedirs(cache_dir, exist_ok=True) # Ensure cache dir exists
    cache_key = get_cache_key(job_text)
    cache_file = os.path.join(cache_dir, f"{cache_key}.json") # Use os.path.join
    if os.path.exists(cache_file):
        try:
            with open(cache_file, "r", encoding='utf-8') as f: # Specify encoding
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"--- WARNING: Failed to load cache file {cache_file}: {e}. Re-fetching. ---")
            return None
    return None


def save_cached_skills(job_text, data):
    cache_dir = "cache"
    os.makedirs(cache_dir, exist_ok=True) # Ensure cache dir exists
    cache_key = get_cache_key(job_text)
    cache_file = os.path.join(cache_dir, f"{cache_key}.json") # Use os.path.join
    try:
        with open(cache_file, "w", encoding='utf-8') as f: # Specify encoding
            json.dump(data, f, indent=2)
    except IOError as e:
        print(f"--- WARNING: Failed to save cache file {cache_file}: {e} ---")
        
def generate_session_id():
    return str(uuid.uuid4())

def save_to_session(session_id, filename, content):
    """Saves content to a session-specific folder using an absolute path."""
    # SESSIONS_BASE_DIR is defined globally near the top
    session_folder = os.path.join(SESSIONS_BASE_DIR, session_id)
    absolute_filepath = os.path.join(session_folder, filename)

    print(f"--- DEBUG [save_to_session]: Attempting to save to absolute path: {absolute_filepath} ---")

    try:
        # Create the session-specific directory if it doesn't exist
        os.makedirs(session_folder, exist_ok=True)

        # Write the file
        with open(absolute_filepath, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"--- SUCCESS [save_to_session]: File saved to {absolute_filepath} ---")
        return absolute_filepath # Return the path where it was saved

    except OSError as e:
        # This is often a Permissions Error!
        print(f"!!! ERROR [save_to_session]: OS Error saving file {absolute_filepath} !!!")
        print(f"!!! ERROR Details: {e} !!!")
        print(f"!!! Check write permissions for the directory: {session_folder} !!!")
        raise # Re-raise the error so the calling function knows it failed

    except Exception as e:
        print(f"!!! ERROR [save_to_session]: Unexpected error saving file {absolute_filepath} !!!")
        print(f"!!! Error type: {type(e).__name__}, Details: {e} !!!")
        traceback.print_exc() # Print full traceback for unexpected errors
        raise # Re-raise the error

def extract_experience_block(resume_text):
    """
    Split the PROFESSIONAL EXPERIENCE section into discrete blocks
    based on the '---' separators (or, as a fallback, on company+date lines).
    """
    lines = resume_text.splitlines()
    blocks = []
    buffer = []

    for line in lines:
        # If we hit a separator line (---), flush the buffer to a block
        if re.match(r'^-{3,}$', line.strip()):
            if buffer:
                blocks.append('\n'.join(buffer).strip())
                buffer = []
        else:
            buffer.append(line)
    # Flush any remaining lines
    if buffer:
        blocks.append('\n'.join(buffer).strip())

    # The first block will be the "PROFESSIONAL EXPERIENCE" header—drop it
    if blocks and blocks[0].strip().upper().startswith("PROFESSIONAL EXPERIENCE"):
        blocks = blocks[1:]

    # Fallback: if somehow everything collapsed into one block, try date-line splitting
    if len(blocks) == 1:
        text = blocks[0]
        lines = text.splitlines()
        boundary_idxs = [
            i for i, line in enumerate(lines)
            if re.match(
                r'^.+\s{2,}(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s+–\s+(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})$',
                line.strip()
            )
        ]
        # include heading as first boundary if found
        heading = [i for i, l in enumerate(lines) if l.strip().startswith("PROFESSIONAL EXPERIENCE")]
        if heading:
            boundary_idxs.insert(0, heading[0])
        if boundary_idxs:
            new_blocks = []
            for start, end in zip(boundary_idxs, boundary_idxs[1:] + [len(lines)]):
                chunk = "\n".join(lines[start:end]).strip()
                if chunk:
                    new_blocks.append(chunk)
            blocks = new_blocks

    return blocks


def get_synonyms(word):
    return [word] # Only return the original word

def match_skills(skills_required, resume_text, scorer=fuzz.token_set_ratio, threshold=85, structured_experience=None):
    matched = []
    missing = []
    scores = {}
    raw_scores = {}

    for skill in skills_required:
        score = scorer(skill.lower(), resume_text.lower())
        scores[skill] = score
        raw_scores[skill] = score
        if score >= threshold:
            matched.append(skill)
        else:
            missing.append(skill)

    return matched, missing, scores, raw_scores

def format_technical_skills_inline(md_text):
    lines = md_text.splitlines()
    new_lines = []
    in_tech_skills = False
    skill_items = []

    for line in lines:
        if line.strip().lower().startswith("technical skills"):
            new_lines.append(line)  # Preserve the header
            in_tech_skills = True
            continue
        if in_tech_skills:
            if line.strip().startswith("- "):
                skill_items.append(line[2:].strip())
            elif not line.strip():
                continue
            else:
                if skill_items:
                    new_lines.append(" | ".join(skill_items))
                    skill_items = []
                in_tech_skills = False
                new_lines.append(line)
        else:
            new_lines.append(line)

    if in_tech_skills and skill_items:
        new_lines.append(" | ".join(skill_items))
    return "\n".join(new_lines)



app = Flask(__name__)
app.register_blueprint(analyze_bp)

CORS(app, resources={r"/*": {"origins": "*"}})

def extract_text_from_docx(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

def call_your_model_api(prompt):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost",
        "X-Title": "Resume Optimizer",
    }

    # 🔍 Step 1: Detect if prompt is a preformatted list of messages (e.g., from hybrid_prompt)
    if isinstance(prompt, list) and all(
            isinstance(msg, dict) and "role" in msg and "content" in msg for msg in prompt):
        messages = prompt
    elif isinstance(prompt, str):  # Handle plain string prompt
        messages = [{"role": "user", "content": prompt}]
    else:
        raise TypeError("Prompt must be a string or a list of message dictionaries.")

    # 🔧 Step 2: Build payload dynamically
    payload = {
        "model": "openai/gpt-3.5-turbo",
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": 3500
    }

    # 🧠 Optional: Print preview of final messages going to LLM
    print("--- DEBUG: Sending payload to LLM ---")

    # 🚀 Step 3: Send to OpenRouter
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            json=payload,
            headers=headers,
            timeout=120  # Add a timeout (e.g., 2 minutes)
        )
        response.raise_for_status()  # Raises HTTPError for bad responses (4xx or 5xx)

        response_data = response.json()

        if "choices" not in response_data or not response_data["choices"]:
            raise ValueError("API response missing 'choices' field or 'choices' is empty.")
        if "message" not in response_data["choices"][0] or "content" not in \
                response_data["choices"][0]["message"]:
            raise ValueError("API response missing 'message' or 'content' in the first choice.")

        content = response_data["choices"][0]["message"]["content"]
        print("--- SUCCESS: Received response from OpenRouter API ---")
        return {"content": content.strip()}  # Strip leading/trailing whitespace
    except requests.exceptions.RequestException as e:
        print(f"!!! ERROR: OpenRouter API request failed: {e} !!!")
        # Consider more specific error handling based on status code if available
        if hasattr(e, 'response') and e.response is not None:
            print(f"!!! Response Status Code: {e.response.status_code} !!!")
            print(f"!!! Response Text: {e.response.text} !!!")
        raise  # Re-raise the exception to be caught by the endpoint handler
    except (ValueError, KeyError, IndexError) as e:
        print(f"!!! ERROR: Invalid response format from OpenRouter API: {e} !!!")
        print(f"!!! Raw Response Data: {response_data if 'response_data' in locals() else 'N/A'} !!!")
        raise  # Re-raise

def detect_resume_sections(resume_text):
    """
    Detects high-level section headers in the resume text.
    Assumes headers are in Title Case or ALL CAPS with no leading bullet/dash.
    Returns a list of section names.
    """
    headers = []
    lines = resume_text.splitlines()

    for line in lines:
        clean = line.strip()
        if (
            clean and len(clean) < 60
            and clean[0].isupper()
            and not clean.startswith(("-", "•", "*"))
        ):
            # Consider title-like lines with no punctuation
            if re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*|[A-Z ]{3,})$', clean):
                headers.append(clean)

    # Optional: deduplicate and preserve order
    seen = set()
    unique_headers = []
    for h in headers:
        if h not in seen:
            seen.add(h)
            unique_headers.append(h)

    return unique_headers



jd_skill_cache = {}



#--- Flask Routes ---




@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    try:  # Start of main try block for the endpoint
        data = request.json
        
        user_tier = data.get("userTier", "unknown")
        free_user_usage_count = data.get("freeUserUsageCount", 0)
        
        print(f"📥 Incoming userTier: {user_tier}")
        print(f"📊 Free usage count: {free_user_usage_count}")
        
        if not data:
            return jsonify({'error': 'Request body must be JSON'}), 400
        
        if user_tier == "free" and free_user_usage_count >= 3:
                    print("🚫 Resume generation blocked for free user – usage limit exceeded")
                    return jsonify({
                        "error": "Free plan limit reached. Upgrade to generate more resumes."
                    }), 403

        resume_raw = data.get("resumeText", "").strip()  # Original resume (potentially base64)
        job_text = data.get("jobText", "").strip()
        extra_answers = data.get("extraAnswers", [])  # Answers to clarification questions
        MAX_ANSWERS_CHARS = 2000
        user_tier = data.get("userTier", "guest")
        if user_tier == "free":
            usage_count = data.get("freeUserUsageCount", 0)
            print(f"📊 [Free Tier] This user has used {usage_count} of 3 allowed resume generations.")

        
        # Get skills from the *previous* analysis step, passed back from frontend
        analyzed_data = data.get("analyzedData") or {}
        extracted_skills = analyzed_data.get("extracted_jd_skills",[])  # Use skills found in analysis
        # Also get missing skills to highlight them in the prompt
        missing_skills_prev = analyzed_data.get("missing_skills", [])
        extra_answers = data.get("extraAnswers", [])  # Answers to clarification questions
        session_id = generate_session_id()
        
        answers_text = "\n".join([f"- {a}" for a in extra_answers if a and a.strip()])[:MAX_ANSWERS_CHARS]     
        
        save_to_session(session_id, "clarification_answers.txt", answers_text)
        print(f"--- DEBUG: Clarification answers saved:\n{answers_text} ---")

        print(f"DEBUG resume_raw={repr(resume_raw[:100])}")
        print(f"DEBUG job_text={repr(job_text[:100])}")

        if not resume_raw or not job_text:
            return jsonify(
                {'error': 'Original resume (resumeText) and job description (jobText) are required'}), 400

        # Decode resume
        try:
            resume_input = decode_resume_base64(resume_raw)
            if not resume_input:
                # This might happen if base64 is invalid or docx extraction fails completely
                print("--- ERROR: Decoded resume text is empty after decode_resume_base64. ---")
                raise ValueError("Resume parsing/decoding resulted in empty text.")
            print("--- INFO: Resume decoded successfully for generation. ---")
        except Exception as e:
            print(f"!!! ERROR: Resume decoding failed in generate: {e} !!!")
            traceback.print_exc()
            return jsonify({"error": f"Resume parsing/decoding failed: {e}"}), 400

        # Extract job title again
        job_title = extract_job_title_from_text(job_text)
        
        # Extract experience blocks BEFORE truncation to preserve full data
        experience_blocks = extract_experience_block(resume_input)  # FIXED LINE
        block_list = "\n\n".join(f"Block {i+1}:\n{blk}" for i, blk in enumerate(experience_blocks))

        # Truncate inputs for the LLM prompt
        MAX_JD_CHARS_GEN = 7000 # Increased slightly if needed, adjust based on model context window
        MAX_RESUME_CHARS_GEN = 7000 # Increased slightly if needed
        job_text_trunc = job_text[:MAX_JD_CHARS_GEN] # Use original job_text for prompt, add title in prompt
        resume_input_trunc = resume_input[:MAX_RESUME_CHARS_GEN]
        
         
        clarification_section = f"""
**Clarification Insights (Candidate Responses):**
{answers_text if answers_text else "None provided."}
"""

        # --- Build the Generation Prompt with forced block+skill integration ---
        missing_skills_bullets = "\n".join(f"- {s}" for s in missing_skills_prev)
       
        generation_prompt = f"""
System:
You are an expert professional resume writer, ATS optimization specialist, and career coach.

User:
**Target Job:** {job_title}

**Job Description (Excerpt):**
{job_text_trunc}

Career Insights:
{answers_text if answers_text else "None provided."}

## Resume Experience Blocks:
(Preserve the original number of experience bullet points per job title. If 4 bullets existed, rewrite 4. If 6, keep 6. Do not reduce count.)
{block_list}

## Missing JD Skills:
{missing_skills_bullets}

## Final Resume Format (write in this structure):
# Executive Summary
# Areas of Expertise
# Key Achievements
# Career Summary
  - Education
  - Certification
# Technical Skills
# Professional Experience

## ATS Optimization Goal:
- You will be evaluated by an ATS system on whether your resume matches the required skills listed above.
- Your objective is to achieve at least **80% skill coverage**.
- Integrate as many of the missing skills as possible into relevant resume sections, using bullet points.
- For each missing skill, insert a **meaningful bullet point** under the most relevant job block.
- Resume must be professional, concise, and ATS-optimized.
- Only submit the final resume output in clean markdown format. Do not include explanations.
"""
        # --- END OF REVISED PROMPT STRING ---


        # --- Retry Loop for LLM Generation if Score < 80 ---
        max_attempts = 5
        attempts = 0
        best_resume_output = ""
        best_score = 0
        tag_validation = []
        resume_output = ""

        while attempts < max_attempts:
            print(f"--- Attempt {attempts+1}/{max_attempts} to generate resume ---")
            try:
                llm_response = call_your_model_api(generation_prompt)
                resume_output = llm_response['content']
                resume_output = re.sub(r"^```(?:[\w\n]+)?\n?", "", resume_output, flags=re.MULTILINE)
                resume_output = re.sub(r"\n?```$", "", resume_output, flags=re.MULTILINE)
                resume_output = resume_output.strip()

                if not resume_output:
                    print(f"⚠️ Empty resume output from LLM on attempt {attempts+1}")
                    attempts += 1
                    continue

                # Score it
                matched, missing, scores, _ = match_skills(
                    extracted_skills, resume_output,
                    scorer=fuzz.partial_ratio, threshold=75
                )
                score = round((len(matched) / len(extracted_skills)) * 100, 2) if extracted_skills else 0
                print(f"✅ ATS Score for attempt {attempts+1}: {score}%")

                if score > best_score:
                    best_resume_output = resume_output
                    best_score = score

                if score >= 80:
                    print(f"🎯 Acceptable score reached: {score}%")
                    break

            except Exception as e:
                print(f"!!! ERROR during resume generation attempt {attempts+1}: {e} !!!")

            attempts += 1

        if not resume_output:
            resume_output = best_resume_output
            print(f"⚠️ Using best available resume with score {best_score}% after {attempts} attempts")

        # --- Post-Generation Analysis ---
        print("--- INFO: Analyzing generated resume for skill matching... ---")

        tag_validation = []

        # --- Added robustness checks here ---
        if not extracted_skills:
             print("--- WARNING: extracted_skills list is empty in generate-resume endpoint. Cannot calculate meaningful score. ---")
             # If skills list was empty from initial analysis, the score will be 0 anyway.
             new_match_score = 0
             new_matched_skills = []
             new_missing_skills = [] # Can't determine missing if extracted is empty
             new_scores = {}

        elif not resume_output.strip(): # Handle case where LLM generated empty/whitespace
            print("--- WARNING: resume_output is empty or whitespace after generation. Score will be 0. ---")
            new_match_score = 0
            new_matched_skills = []
            new_missing_skills = list(extracted_skills) # All are missing if resume is empty
            new_scores = {}
        else:
             # Recompute best resume match metrics
            resume_output = best_resume_output
            saved_path = save_to_session(session_id, "resume_optimized.txt", resume_output)
            print(f"--- INFO: Best Resume Output Score: {best_score}% ---")

            # Match skills from best resume
            new_match_score = best_score
            new_matched_skills, new_missing_skills, new_scores, _ = match_skills(
                extracted_skills,
                resume_output,
                scorer=fuzz.partial_ratio,
                threshold=75
            )

            # Calculate score only if extracted_skills is not empty (handled by the outer elif structure)
            new_match_score = round((len(new_matched_skills) / len(extracted_skills)) * 100, 2)

            # Step 2: Log & store tag validation
            tagged_skills = [f"**[{s.strip()}]**" for s in missing_skills_prev if s.strip()]
            tag_validation = []
            for tag in tagged_skills:
                found = tag.lower() in best_resume_output.lower()
                tag_validation.append({"tag": tag, "present": found})
                if not found:
                    print(" ⚠️  Missing Tag:", tag)
            if all(t['present'] for t in tag_validation):
                print("--- SUCCESS: All missing skills were successfully inserted as tags. ---")
        
        # --- End robustness checks ---


        # --- ADD DEBUG PRINT STATEMENTS HERE ---
        print(f"--- DEBUG [Generate]: Calculated new_match_score: {new_match_score} ---")
        print(f"--- DEBUG [Generate]: extracted_skills count: {len(extracted_skills)} ---")
        # Printing the whole analytics object can be verbose, maybe shorten or only include key parts
        # print(f"--- DEBUG [Generate]: analytics object being returned: {json.dumps({
        #     "match_score": new_match_score,
        #     "matched_skills": new_matched_skills,
        #     "missing_skills": new_missing_skills,
        #     "skill_scores": new_scores,
        #     "extracted_jd_skills": extracted_skills, # Be mindful of size if this list is huge
        # }, indent=2)} ---")
        print(f"--- DEBUG [Generate]: analytics object being returned (partial): match_score={new_match_score}, matched_skills_count={len(new_matched_skills)}, missing_skills_count={len(new_missing_skills)} ---")
        # --- END DEBUG PRINT STATEMENTS ---

        session_data = {
            "resume": resume_output,
            "summary": "",
            "resume_raw": resume_raw,
            "resume_decoded": resume_input,
            "experience_blocks": experience_blocks,       
            "jd_text": job_text,
            "jd_extracted_skills": extracted_skills,
            "analytics": {
                "match_score": new_match_score,
                "matched_skills": new_matched_skills,
                "missing_skills": new_missing_skills,
                "skill_scores": new_scores,
                "extracted_jd_skills": extracted_skills,
                "tag_validation": tag_validation
            },
            "session_id": session_id
        }

        print("--- DEBUG: Session data to be saved ---")
        print(json.dumps(session_data, indent=2))


        base_dir = os.path.dirname(os.path.dirname(__file__))  # project root
        session_path = os.path.join(base_dir, 'sessions', f'{session_id}.json')
        os.makedirs(os.path.dirname(session_path), exist_ok=True)

        with open(session_path, 'w') as f:
            json.dump(session_data, f)
            print(f"✅ Session saved to: {session_path}")

        
        
        # --- Response ---
        return jsonify({
            "resume": resume_output,  # The generated resume text
            "analytics": {
                "match_score": new_match_score,  # <--- This is the score sent
                "matched_skills": new_matched_skills,
                "missing_skills": new_missing_skills,
                "skill_scores": new_scores,
                "extracted_jd_skills": extracted_skills,
            },
            "session_id": session_id,
            "confidence_flags": [] # Include if implemented
        })

    # --- Enhanced Catch-all Exception Handler for /generate-resume ---
    except Exception as e:
        # Catch-all for unexpected errors during request processing
        print(f"!!! FATAL ERROR in /generate-resume endpoint !!!")
        print(f"Error type: {type(e).__name__}")
        print(f"Error details: {e}")
        traceback.print_exc()  # Prints the full traceback to Flask console
        # Return a 500 Internal Server Error response
        return jsonify({'error': f'Failed to generate resume due to an unexpected server error: {str(e)}'}), 500

@app.route('/session/<session_id>', methods=['GET'])
def fetch_session_data(session_id):
    try:
        base_sessions_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'sessions')
        session_file_path = os.path.join(base_sessions_dir, f"{session_id}.json")

        print(f"🔍 Looking for session file at: {session_file_path}")

        with open(session_file_path, 'r') as f:
            data = json.load(f)

        # Use base_sessions_dir here to keep it clean
        session_subdir = os.path.join(base_sessions_dir, session_id)
        analysis_path = os.path.join(session_subdir, "ai_analysis.txt")

        if os.path.exists(analysis_path):
            try:
                with open(analysis_path, 'r', encoding='utf-8') as f:
                    full_analysis_text = f.read()
                if "**Clarification Questions:**" in full_analysis_text:
                    summary_only = full_analysis_text.split("**Clarification Questions:**")[0].strip()
                    data["summary"] = summary_only
            except Exception as e:
                print(f"⚠️ Could not parse ai_analysis.txt for session {session_id}: {e}")

        return jsonify(data)

    except FileNotFoundError:
        print("❌ Session file not found.")
        return jsonify({"error": "Session not found"}), 404
    except Exception as e:
        print(f"❌ Unexpected error in fetch_session_data: {e}")
        return jsonify({"error": str(e)}), 500



@app.route('/download-resume/<session_id>', methods=['GET'])
def download_resume(session_id):
    if not session_id or not re.match(r'^[a-zA-Z0-9-]+$', session_id):
        return jsonify({"error": "Invalid session ID format"}), 400

    session_folder = os.path.join(SESSIONS_BASE_DIR, session_id)
    resume_txt_path = os.path.join(session_folder, "resume_optimized.txt")
    pdf_output_path = os.path.join(session_folder, "resume_optimized.pdf")

    if not os.path.exists(resume_txt_path):
        return jsonify({"error": "Optimized resume text file not found for this session"}), 404

    try:
        with open(resume_txt_path, "r", encoding="utf-8") as f:
            raw_resume_text = f.read()
            
            # Convert Markdown bullets to proper spacing
            cleaned_md = re.sub(r'(?<!\n)- ', '\n- ', raw_resume_text)

            # Convert Technical Skills bullets to pipe-separated string
            formatted_md = format_technical_skills_inline(cleaned_md)


        if not raw_resume_text.strip():
            return jsonify({"error": "Resume content is empty, cannot generate PDF."}), 400

        # Convert markdown to styled HTML using markdown2
        # Ensure bullet points stay on separate lines
        cleaned_md = re.sub(r'(?<!\n)- ', '\n- ', raw_resume_text)  # Ensure each "- " starts a new line
        formatted_md = format_technical_skills_inline(cleaned_md)
        converted_html = markdown2.markdown(formatted_md)
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset='UTF-8'>
            <title>Optimized Resume</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }}
                h1, h2, h3, h4 {{ color: #333; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
                ul {{ margin-top: 0; }}
                strong {{ color: #000; }}
                hr {{ border: none; border-top: 1px solid #ccc; margin: 1em 0; }}
            </style>
        </head>
        <body>
            {converted_html}
        </body>
        </html>
        """

        if not PDF_CONFIG:
            return send_file(resume_txt_path, as_attachment=True, download_name="optimized_resume.txt", mimetype='text/plain')

        pdfkit.from_string(full_html, pdf_output_path, configuration=PDF_CONFIG, options={'enable-local-file-access': None})

        if os.path.exists(pdf_output_path):
            return send_file(
                pdf_output_path,
                as_attachment=True,
                download_name=f"Optimized_Resume_{session_id[:8]}.pdf",
                mimetype='application/pdf'
            )
        else:
            return jsonify({"error": "PDF file was not created as expected."}), 500

    except Exception as e:
        print(f"!!! ERROR in /download-resume: {e} !!!")
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500

@app.route("/session/<session_id>/data", methods=["GET"])
def get_full_session_data(session_id):
    try:
        folder = os.path.join(SESSIONS_BASE_DIR, session_id)

        def safe_read(filename):
            path = os.path.join(folder, filename)
            return open(path, "r", encoding="utf-8").read().strip() if os.path.exists(path) else ""

        resume = safe_read("resume_optimized.txt")
        jd = safe_read("jd.txt")
        decoded_resume = safe_read("resume_decoded.txt")
        analytics = json.loads(safe_read("resume_analytics.json")) if os.path.exists(os.path.join(folder, "resume_analytics.json")) else {}

        return jsonify({
            "session_id": session_id,
            "resume": resume,
            "jobText": jd,
            "resumeText": decoded_resume,
            "analytics": analytics
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500



# --- Main Execution ---
if __name__ == '__main__':
    # Ensure cache directory exists on startup
    os.makedirs("cache", exist_ok=True)
    # Ensure base sessions directory exists on startup (optional, makedirs handles it later too)
    try:
         os.makedirs(SESSIONS_BASE_DIR, exist_ok=True)
         print(f"--- INFO: Ensured sessions base directory exists: {SESSIONS_BASE_DIR} ---")
    except OSError as e:
         print(f"!!! WARNING: Could not create sessions base directory on startup: {e}. File saving might fail later if permissions are incorrect. !!!")

    print("--- Starting Flask Development Server ---")
    # Use host='0.0.0.0' to make it accessible on your network if needed
    app.run(debug=True, host='127.0.0.1', port=5000)