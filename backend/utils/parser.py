import re
import os
import base64
import tempfile
import fitz # PyMuPDF
from docx import Document
from typing import List

KNOWN_SKILLS = {
    "python", "java", "javascript", "typescript", "c++", "c#", "sql", "mongodb",
    "aws", "azure", "gcp", "docker", "kubernetes", "git", "jira", "confluence",
    "html", "css", "react", "node.js", "django", "flask",
    "excel", "power bi", "tableau", "salesforce", "dynamics 365",
    "sap", "notion", "miro", "figma", "linux", "windows", "agile", "scrum"
}


def extract_pdf_text(path: str) -> str:
    doc: fitz.Document = fitz.open(path)
    return "\n".join([page.get_text() for page in doc])  # type: ignore

def extract_docx_text(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_resume_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_pdf_text(path)
    elif ext == '.docx':
        return extract_docx_text(path)
    else:
        raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")

def parse_resume_fields(text):
    sections = split_into_sections(text)
    experience = extract_experience_block(text)

    return {
        "summary": sections.get("summary", "")[:500] + "...",
        "experience": experience,
        "skills": extract_skills(text),
        "education": extract_education_block(text),
        "raw_text": text[:1000] 
    }


def extract_experience_block(resume_text: str) -> list:
    sections = split_into_sections(resume_text)
    exp_text = ""

    for key in sections:
        if 'experience' in key.lower() or 'career summary' in key.lower():
            exp_text = sections[key]
            break

    if not exp_text:
        return []

    lines = [line.strip() for line in exp_text.splitlines() if line.strip()]
    jobs = []
    current_job = None

    job_header_pattern = re.compile(
        r'^(?P<title>[^\-–—|@]+)\s*[-–—|@]\s*(?P<company>[^()|]+)\s*[()|]?\s*(?P<period>(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}.*)?',
        re.IGNORECASE
    )

    for line in lines:
        match = job_header_pattern.match(line)
        if match:
            if current_job:
                jobs.append(current_job)

            current_job = {
                "title": match.group("title").strip(),
                "company": match.group("company").strip(),
                "period": match.group("period").strip() if match.group("period") else "",
                "responsibilities": ""
            }
        elif current_job:
            current_job["responsibilities"] += line + "\n"

    if current_job:
        jobs.append(current_job)

    return jobs





def extract_skills(text: str) -> List[str]:
    text_lower = text.lower()
    return sorted({skill for skill in KNOWN_SKILLS if skill in text_lower})


def split_into_sections(text):
    """Split resume text into sections using common headers."""
    headers = [
        "contact information", "contact", "profile", "summary", "about me", "objective",
        "experience", "work experience", "employment history", "professional experience",
        "education", "academic background", "skills", "technical skills",
        "certifications", "achievements", "projects", "languages"
    ]
    pattern = "|".join(rf"{re.escape(header)}" for header in headers)
    matches = list(re.finditer(pattern, text, flags=re.IGNORECASE))  # ✅ use flags, not inline
    sections = {}
    
    for i, match in enumerate(matches):
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        section_title = match.group().lower()
        sections[section_title] = text[start:end].strip()

    return sections

def extract_education_block(text: str) -> List[dict]:
    sections = split_into_sections(text)
    edu_text = ""

    for key in sections:
        if 'education' in key.lower() or 'certification' in key.lower():
            edu_text = sections[key]
            break

    if not edu_text:
        return []

    lines = [line.strip() for line in edu_text.splitlines() if line.strip()]
    education_entries = []

    for line in lines:
        # Extract completion date like "May 2025" or just "2025"
        date_match = re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}", line, re.IGNORECASE)
        completed = date_match.group(0).strip() if date_match else ""

        # Remove date from line
        clean_line = line.replace(completed, "").strip() if completed else line

        # Try to split by |, ,, or -
        parts = [part.strip() for part in re.split(r"\||,|-", clean_line)]
        course = parts[0] if len(parts) > 0 else ""
        school = parts[1] if len(parts) > 1 else ""

        level_guess = "Bachelor" if "bachelor" in course.lower() else (
                      "Diploma" if "diploma" in course.lower() else (
                      "Certificate" if "cert" in course.lower() else ""))

        education_entries.append({
            "school": school,
            "course": course,
            "level": level_guess,
            "completed": completed
        })

    return education_entries

def decode_resume_base64(resume_raw):
    """Helper to decode a base64-encoded .docx file or plain text."""
    if resume_raw.startswith('UEsDB'):
        try:
            docx_bytes = base64.b64decode(resume_raw)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                f.write(docx_bytes)
                path = f.name
            doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            raise ValueError(f"Resume decode failed: {e}")
    else:
        return resume_raw


def extract_job_title_from_text(jd_text):
    """
    Attempts to extract a meaningful job title from job description text.
    """
    # Try to match common title indicators
    match = re.search(r'(?i)(position|role|job title)[:\-]?\s*(.+)', jd_text)
    if match:
        return match.group(2).strip()

    # Fallback: first non-empty, non-generic line
    lines = jd_text.strip().split('\n')
    for line in lines:
        cleaned = line.strip()
        if cleaned and not cleaned.lower().startswith(("about", "who we", "join", "our company", "the company")):
            return cleaned
    return "Untitled Role"